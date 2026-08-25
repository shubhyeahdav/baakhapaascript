import io
import math
import os
import textwrap
import time
import urllib.request as urllib_request
from datetime import datetime
from typing import Optional

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Pt

import scene_sync
import screenplay

# ---------------------------------------------------------------------------
# Devanagari support
#
# ReportLab's built-in Courier has no Devanagari glyphs, so Nepali dialogue
# rendered with it comes out as blank boxes. We register a Unicode TTF instead
# and fall back to Courier only if none is found.
#
# For deployment, bundle Noto Sans Devanagari (SIL Open Font License, so it is
# redistributable) at assets/NotoSansDevanagari-Regular.ttf. The Windows path
# below is a local-development convenience only — Nirmala UI is a Microsoft
# font and must not be shipped.
# ---------------------------------------------------------------------------
_ASSETS = os.path.join(os.path.dirname(__file__), "assets")

# (path, subfont_index) — .ttc collections need an index; .ttf files ignore it.
_FONT_CANDIDATES = [
    (os.path.join(_ASSETS, "NotoSansDevanagari-Regular.ttf"), None),
    (os.path.join(_ASSETS, "NotoSansDevanagari.ttf"), None),
    ("/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf", None),
    (r"C:\Windows\Fonts\Nirmala.ttc", 0),  # local dev only — not redistributable
]

_BOLD_CANDIDATES = [
    (os.path.join(_ASSETS, "NotoSansDevanagari-Bold.ttf"), None),
    ("/usr/share/fonts/truetype/noto/NotoSansDevanagari-Bold.ttf", None),
    (r"C:\Windows\Fonts\NirmalaB.ttf", None),
]


def _try_register(name: str, path: str, subfont) -> bool:
    if not os.path.exists(path):
        return False
    try:
        if subfont is None:
            pdfmetrics.registerFont(TTFont(name, path))
        else:
            pdfmetrics.registerFont(TTFont(name, path, subfontIndex=subfont))
        return True
    except Exception:
        return False


def _register_devanagari_font():
    """Register a Devanagari-capable font.

    Returns (body_font, bold_font, resolved_path). The path is reported so the
    production preflight can tell a bundled OFL face apart from Windows' Nirmala
    — both make `DEVANAGARI_READY` true, but only one is legal to deploy.
    """
    body, bold = "Courier", "Courier-Bold"
    resolved = None

    for path, subfont in _FONT_CANDIDATES:
        if _try_register("BaakhapaaScript", path, subfont):
            body = "BaakhapaaScript"
            resolved = path
            break

    if body == "Courier":
        print(
            "WARNING: no Devanagari font found — PDF export will render Nepali "
            "text as blank boxes. Add NotoSansDevanagari-Regular.ttf to "
            "baakhapaa-backend/assets/."
        )
        return body, bold, None

    for path, subfont in _BOLD_CANDIDATES:
        if _try_register("BaakhapaaScript-Bold", path, subfont):
            bold = "BaakhapaaScript-Bold"
            break
    if bold == "Courier-Bold":
        bold = body  # better an unbolded Devanagari heading than a missing glyph

    return body, bold, resolved


BODY_FONT, BOLD_FONT, RESOLVED_FONT_PATH = _register_devanagari_font()
DEVANAGARI_READY = BODY_FONT != "Courier"


def _has_devanagari(text: str) -> bool:
    return any("ऀ" <= ch <= "ॿ" for ch in text)


def _font_for(line: str) -> str:
    """Courier for Latin lines, the Unicode font only where it's needed.

    Screenplay format depends on a monospaced page, and the available
    Devanagari faces are proportional. Switching per line keeps English
    action and sluglines correctly monospaced while still rendering Nepali
    dialogue instead of blank boxes.
    """
    return BODY_FONT if (DEVANAGARI_READY and _has_devanagari(line)) else "Courier"


# ---------------------------------------------------------------------------
# Screenplay page layout
#
# US Letter, not A4: a screenplay is a US Letter document everywhere in the
# world that buys one, and this printed A4 with every line starting at 1.25"
# and hard-cut at 90 characters -- which on a 595pt page ran a third of every
# long action line off the right edge of the paper before the cut even applied.
#
# The row model lives in `screenplay.layout_rows()` so the page the writer is
# told they are on and the page this prints are the same page.
# ---------------------------------------------------------------------------
PAGE_W, PAGE_H = letter
TOP_MARGIN = 72.0
BOTTOM_MARGIN = 72.0
FONT_SIZE = 12
# 45 rows between one-inch margins. Derived rather than chosen, so the geometry
# and `screenplay.PAGE_LINES` cannot drift apart.
LEADING = (PAGE_H - TOP_MARGIN - BOTTOM_MARGIN) / screenplay.PAGE_LINES

# Left edge of each element, in points from the page edge. The standard
# screenplay measures: action at 1.5", dialogue at 2.5", parenthetical at 3.1",
# character cue at 3.7". Every line used to be drawn flush at 1.25", which is
# the difference between a screenplay and a text file with a title page.
INDENT = {
    "scene_heading": 108.0,
    "action": 108.0,
    "character": 266.0,
    "parenthetical": 223.0,
    "dialogue": 180.0,
    "transition": 108.0,
    "blank": 108.0,
}


def _draw_screenplay_pages(c, script_content: str) -> int:
    """Draw the screenplay, one printed page per sheet. Returns pages drawn.

    Assumes the caller has already put a fresh page under the pen, and leaves
    the last page open for the caller to close.
    """
    rows = screenplay.layout_rows(script_content) or [
        screenplay.Row("", "blank", None, 1)
    ]

    # Scene numbers, in both margins beside the slugline. A production script
    # without them cannot be broken down, scheduled or shot-listed by anyone.
    scene_number = {}
    seen = 0
    for row in rows:
        if row.type == "scene_heading" and not row.continued:
            seen += 1
            scene_number[row.source_line] = seen

    per_page = screenplay.PAGE_LINES
    pages = max(1, math.ceil(len(rows) / per_page))

    for page_index in range(pages):
        chunk = rows[page_index * per_page:(page_index + 1) * per_page]

        # Page number, top right, the printed convention. Page one carries none.
        if page_index:
            c.setFont("Courier", FONT_SIZE)
            c.drawRightString(PAGE_W - 72, PAGE_H - 50, f"{page_index + 1}.")

        # A page that opens mid-speech repeats the cue as SPEAKER (CONT'D).
        # Drawn in the top margin rather than as a row, because taking a row
        # would push the page count away from the one the editor shows.
        first = chunk[0]
        prev = rows[page_index * per_page - 1] if page_index else None
        speaker = None
        if first.type in ("dialogue", "parenthetical") and first.character:
            speaker = first.character
        elif prev is not None and prev.type == "character":
            speaker = prev.text
        if page_index and speaker:
            c.setFont(_font_for(speaker), FONT_SIZE)
            c.drawString(INDENT["character"], PAGE_H - TOP_MARGIN + LEADING,
                         f"{speaker} (CONT'D)")

        y = PAGE_H - TOP_MARGIN
        for row in chunk:
            if row.text:
                c.setFont(_font_for(row.text), FONT_SIZE)
                if row.type == "scene_heading" and not row.continued:
                    number = scene_number.get(row.source_line)
                    if number:
                        c.drawString(INDENT["scene_heading"] - 36, y, str(number))
                        c.drawString(PAGE_W - 72 + 12, y, str(number))
                c.drawString(INDENT.get(row.type, INDENT["action"]), y, row.text)
            y -= LEADING

        # (MORE) under the last row when the speech continues overleaf.
        last = chunk[-1]
        following = rows[(page_index + 1) * per_page] if (page_index + 1) * per_page < len(rows) else None
        if following is not None:
            carries = (
                (following.type in ("dialogue", "parenthetical") and following.character)
                or last.type == "character"
            )
            if carries:
                c.setFont("Courier", FONT_SIZE)
                c.drawString(INDENT["dialogue"], y, "(MORE)")

        if page_index < pages - 1:
            c.showPage()

    return pages


def export_script_pdf(script_content: str, title: str = "Untitled") -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)

    c.setFont(BOLD_FONT if _has_devanagari(title) else "Courier-Bold", 16)
    c.drawCentredString(PAGE_W / 2, PAGE_H - 100, title)
    c.showPage()

    _draw_screenplay_pages(c, script_content)
    c.save()
    buffer.seek(0)
    return buffer.read()


_FDX_TYPE = {
    "scene_heading": "Scene Heading",
    "action": "Action",
    "character": "Character",
    "parenthetical": "Parenthetical",
    "dialogue": "Dialogue",
    "transition": "Transition",
    # Final Draft has a Shot type; it has nothing for a montage marker or a
    # television act break, which are conventionally written as action and
    # read correctly by every tool when they are.
    "shot": "Shot",
    "montage": "Action",
    "act_break": "Action",
}

# Anything the parser learns to recognise later still has to export. A KeyError
# here would take out the whole .fdx download over one unfamiliar line, and
# Action is always a safe home for a line whose type is not known.
_FDX_FALLBACK = "Action"


def export_script_fdx(script_content: str, title: str = "Untitled") -> bytes:
    """Final Draft (.fdx) export.

    The format every other screenwriting tool reads and writes. Without it a
    writer cannot hand the script to anyone working in Final Draft, Celtx, or
    Arc Studio without retyping it — PDF and Word are both read-only as far as
    screenplay structure is concerned.

    .fdx is plain XML: a flat list of typed paragraphs. The element types come
    straight from the parser, which is why this is short.
    """
    import xml.etree.ElementTree as ET

    import screenplay

    doc = ET.Element("FinalDraft", {
        "DocumentType": "Script", "Template": "No", "Version": "5",
    })
    content = ET.SubElement(doc, "Content")

    for el in screenplay.parse(script_content):
        para = ET.SubElement(content, "Paragraph", {"Type": _FDX_TYPE.get(el.type, _FDX_FALLBACK)})
        ET.SubElement(para, "Text").text = el.text

    title_page = ET.SubElement(doc, "TitlePage")
    tp_content = ET.SubElement(title_page, "Content")
    para = ET.SubElement(tp_content, "Paragraph", {"Alignment": "Center"})
    ET.SubElement(para, "Text").text = title

    return b'<?xml version="1.0" encoding="UTF-8" standalone="no"?>\n' + ET.tostring(
        doc, encoding="utf-8"
    )


def export_script_word(script_content: str, title: str = "Untitled") -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in script_content.split("\n"):
        p = doc.add_paragraph(line)
        p.style.font.name = "Courier New"
        p.style.font.size = Pt(12)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ---------------------------------------------------------------------------
# Production package
#
# PRD US5 promises "script + storyboard + shot list in one click", handed
# straight to a production team. What this produced for a long time was a title
# page, the script, and a page reading `Scene 1 | Wide Shot` — no locations, no
# cast, no images. Nobody can shoot from that.
#
# Everything below assembles the document from the scene rows, which
# `scene_sync` keeps in step with the written draft, so the shot list describes
# the screenplay as it stands rather than the outline it started as.
# ---------------------------------------------------------------------------

# Storyboard images live at a URL (DALL-E, or a placeholder host in demo mode),
# so putting them in the PDF means fetching them. That is a network call inside
# an export, and it must never be why a download hangs or fails:
#   * each fetch is individually timed out,
#   * the whole set shares a budget,
#   * any failure degrades to a captioned frame box, and
#   * it can be switched off entirely with one env var.
# DALL-E URLs also expire after about an hour, so the fallback is a normal
# outcome for an old board rather than an error.
EMBED_STORYBOARD_IMAGES = os.getenv("EMBED_STORYBOARD_IMAGES", "true").lower() != "false"
IMAGE_FETCH_TIMEOUT = float(os.getenv("STORYBOARD_IMAGE_TIMEOUT", "5"))
IMAGE_FETCH_BUDGET = float(os.getenv("STORYBOARD_IMAGE_BUDGET", "20"))
MAX_IMAGE_BYTES = 8 * 1024 * 1024


def _is_public_host(host: str) -> bool:
    """True only if every address `host` resolves to is on the public internet.

    An export is the one place this server fetches a URL, which makes it the one
    place an attacker could aim it inward — at cloud metadata (169.254.169.254),
    at localhost admin ports, at anything on the private network. Frames are
    written by the server, so the URLs should always be external anyway; this is
    the second lock, for a redirect or a provider hostname that later resolves
    somewhere unexpected.
    """
    import ipaddress
    import socket

    if not host:
        return False
    try:
        infos = socket.getaddrinfo(host, None)
    except OSError:
        return False

    for info in infos:
        try:
            ip = ipaddress.ip_address(info[4][0])
        except ValueError:
            return False
        if (ip.is_private or ip.is_loopback or ip.is_link_local
                or ip.is_reserved or ip.is_multicast or ip.is_unspecified):
            return False
    return bool(infos)


class _NoRedirects(urllib_request.HTTPRedirectHandler):
    """Refuse redirects outright.

    Validating the host and then following a redirect validates nothing: the
    redirect target is chosen after the check. Storyboard hosts serve images
    directly, so refusing costs nothing real.
    """

    def redirect_request(self, req, fp, code, msg, headers, newurl):
        return None


_SAFE_OPENER = urllib_request.build_opener(_NoRedirects)


def _fetch_image(url: str, deadline: float):
    """Return an ImageReader for `url`, or None if it cannot be had safely and
    cheaply. Never raises — a missing image is a degraded frame, not a failure."""
    if not EMBED_STORYBOARD_IMAGES or not url:
        return None
    if time.monotonic() >= deadline:
        return None

    # A frame stored inline. The gpt-image models return base64 rather than a
    # hosted URL, so this is now the normal case for a real board — and unlike
    # the old expiring DALL·E links, it still embeds a week later. No network,
    # so no timeout and no host check apply.
    if url.startswith("data:image/"):
        try:
            import base64

            from reportlab.lib.utils import ImageReader

            raw = base64.b64decode(url.split(",", 1)[1])
            if not raw or len(raw) > MAX_IMAGE_BYTES:
                return None
            return ImageReader(io.BytesIO(raw))
        except Exception:
            return None

    if not url.startswith(("http://", "https://")):
        return None
    try:
        from urllib.parse import urlparse

        from reportlab.lib.utils import ImageReader

        parsed = urlparse(url)
        if not _is_public_host(parsed.hostname):
            print(f"Storyboard image host refused (not public): {parsed.hostname}")
            return None

        remaining = min(IMAGE_FETCH_TIMEOUT, deadline - time.monotonic())
        with _SAFE_OPENER.open(url, timeout=remaining) as resp:
            raw = resp.read(MAX_IMAGE_BYTES + 1)
        if not raw or len(raw) > MAX_IMAGE_BYTES:
            return None
        return ImageReader(io.BytesIO(raw))
    except Exception:
        return None


def _slugline(scene: dict, fallback: str) -> str:
    """The scene as a first AD would write it on a call sheet."""
    draft = scene_sync.read_draft(scene)
    if draft.get("heading"):
        return draft["heading"]

    location = (scene.get("location") or "").strip()
    if not location:
        return (scene.get("title") or fallback).upper()

    prefix = "INT." if draft.get("interior", True) else "EXT."
    when = (draft.get("time_of_day") or "").strip()
    return f"{prefix} {location.upper()}" + (f" - {when.upper()}" if when else "")


def _shot_rows(frames: list, scenes: list) -> list:
    """Join frames to their scenes into the rows a shot list is made of."""
    scenes_by_id = {s["id"]: s for s in (scenes or []) if s.get("id")}
    ordered = sorted(frames or [], key=lambda f: f.get("order_index") or 0)

    rows = []
    for i, frame in enumerate(ordered, start=1):
        scene = scenes_by_id.get(frame.get("scene_id")) or {}
        draft = scene_sync.read_draft(scene)
        cast = draft.get("characters") or scene_sync.read_characters(scene)
        rows.append({
            "number": i,
            "slugline": _slugline(scene, f"Scene {i}"),
            "shot_type": frame.get("shot_type") or "-",
            "cast": [str(c) for c in cast],
            "action": (draft.get("summary") or scene.get("description") or "").strip(),
            "camera_notes": (frame.get("camera_notes") or "").strip(),
            "image_url": frame.get("image_url") or "",
            "emotional_beat": (scene.get("emotional_beat") or "").strip(),
        })
    return rows


def _wrap(text: str, width: int) -> list:
    return textwrap.wrap(text, width=width) or [""]


def export_production_package(script_content: str, frames: list, title: str = "Untitled",
                              scenes: Optional[list] = None) -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    rows = _shot_rows(frames, scenes)

    def new_page(heading=None):
        c.showPage()
        y = height - 72
        if heading:
            c.setFont("Courier-Bold", 14)
            c.drawString(72, y, heading)
            y -= 28
        c.setFont("Courier", 10)
        return y

    def draw_wrapped(lines, x, y, size, leading, bottom, heading):
        for line in lines:
            if y < bottom:
                y = new_page(heading)
            c.setFont(_font_for(line), size)
            c.drawString(x, y, line)
            y -= leading
        return y

    # --- title page --------------------------------------------------------
    c.setFont(BOLD_FONT if _has_devanagari(title) else "Courier-Bold", 20)
    c.drawCentredString(width / 2, height / 2 + 20, title[:60])
    c.setFont("Courier", 11)
    c.drawCentredString(width / 2, height / 2 - 12, "PRODUCTION PACKAGE")
    c.setFont("Courier", 9)
    shots = f"{len(rows)} shot" + ("s" if len(rows) != 1 else "")
    c.drawCentredString(
        width / 2, height / 2 - 40,
        f"Script  -  {shots}  -  Generated {datetime.now().strftime('%Y-%m-%d')}",
    )

    # --- screenplay --------------------------------------------------------
    # The same layout the standalone PDF export uses. It had its own copy of
    # this loop, breaking pages on y-geometry rather than on the row count, so
    # the screenplay inside a production package paginated differently from the
    # screenplay the writer had just exported -- and truncated at 90 characters
    # in exactly the same way.
    c.showPage()
    _draw_screenplay_pages(c, script_content)

    # --- shot list ---------------------------------------------------------
    # The document a crew actually carries. One block per shot, so a long
    # slugline or note wraps instead of being cut off at the page edge.
    y = new_page("SHOT LIST")
    if not rows:
        c.setFont("Courier", 10)
        c.drawString(72, y, "No storyboard frames yet - generate a storyboard to fill this in.")

    for row in rows:
        if y < 130:
            y = new_page("SHOT LIST (cont.)")

        c.setFont("Courier-Bold", 10)
        c.drawString(72, y, "%02d" % row["number"])
        c.drawRightString(width - 72, y, row["shot_type"].upper())
        c.setFont(_font_for(row["slugline"]), 10)
        c.drawString(100, y, row["slugline"][:58])
        y -= 15

        c.setFont("Courier", 9)
        if row["cast"]:
            c.drawString(100, y, ("Cast: " + ", ".join(row["cast"]))[:76])
            y -= 13
        if row["emotional_beat"]:
            c.drawString(100, y, ("Beat: " + row["emotional_beat"])[:76])
            y -= 13
        if row["action"]:
            y = draw_wrapped(_wrap(row["action"], 74), 100, y, 9, 12, 70, "SHOT LIST (cont.)")
        if row["camera_notes"]:
            y = draw_wrapped(_wrap("Camera: " + row["camera_notes"], 74), 100, y, 9, 12, 70,
                             "SHOT LIST (cont.)")

        y -= 10
        c.setStrokeColorRGB(0.8, 0.8, 0.8)
        c.line(72, y, width - 72, y)
        y -= 16

    # --- storyboard --------------------------------------------------------
    # Two frames a page: large enough to read a composition from, small enough
    # that a 24-frame board is 12 pages rather than 24.
    if rows:
        deadline = time.monotonic() + IMAGE_FETCH_BUDGET
        frame_w = width - 144
        frame_h = frame_w * 9 / 16
        y = new_page("STORYBOARD")

        for row in rows:
            if y - frame_h - 46 < 60:
                y = new_page("STORYBOARD (cont.)")

            top = y - frame_h
            image = _fetch_image(row["image_url"], deadline)
            if image is not None:
                try:
                    c.drawImage(image, 72, top, width=frame_w, height=frame_h,
                                preserveAspectRatio=True, anchor="c", mask="auto")
                except Exception:
                    image = None
            if image is None:
                c.setStrokeColorRGB(0.75, 0.75, 0.75)
                c.rect(72, top, frame_w, frame_h)
                c.setFont("Courier", 9)
                c.setFillColorRGB(0.45, 0.45, 0.45)
                c.drawCentredString(width / 2, top + frame_h / 2,
                                    "[ " + row["shot_type"] + " - frame image not embedded ]")
                c.setFillColorRGB(0, 0, 0)

            y = top - 16
            c.setFont("Courier-Bold", 10)
            c.drawString(72, y, "%02d  %s" % (row["number"], row["shot_type"].upper()))
            y -= 13
            c.setFont(_font_for(row["slugline"]), 9)
            c.drawString(72, y, row["slugline"][:78])
            y -= 24

    c.save()
    buffer.seek(0)
    return buffer.read()
